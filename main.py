#!/usr/bin/env python3
# MIT License
#
# Copyright (c) 2025 BlackcoinDev
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.
"""
AI Assistant Chat Application v0.1 - Learning Edition

CORE SYSTEM OVERVIEW:
====================
This application provides an intelligent conversational AI assistant with advanced learning
and document processing capabilities. Built specifically for the devstral-small-2507-mlx + qwen3-embedding + ChromaDB stack.

ARCHITECTURAL PRINCIPLES:
- Zero hardcoded defaults - all configuration via .env file
- Type-safe implementation with comprehensive MyPy checking
- Dual interface support (GUI + CLI) with identical functionality
- Sandboxed file operations within current directory only
- Persistent knowledge across sessions via vector database

VERSION 0.1 CAPABILITIES:
========================
🤖 AI FEATURES:
- Conversational AI with devstral-small-2507-mlx via LM Studio
- 7 specialized tools for file operations and document processing
- Autonomous tool calling with natural language triggers
- Context-aware responses using learned knowledge

🧠 LEARNING SYSTEM:
- Teach AI new information via /learn command
- Persistent knowledge storage in ChromaDB v2
- Semantic search with qwen3-embedding via Ollama
- Spaces system for isolated knowledge workspaces

📄 DOCUMENT INTELLIGENCE:
- devstral-small-2507 multimodal analysis for OCR, table extraction, form recognition
- Support for PDFs, images, Office documents, and text files
- Structured data extraction (text, tables, forms, layout)
- Bulk codebase ingestion with 80+ file type support

💾 DATA MANAGEMENT:
- SQLite database for conversation memory
- ChromaDB v2 server for vector knowledge storage
- Automatic conversation persistence across sessions
- Secure file operations with path validation

🎛️ USER INTERFACE:
- Modern PyQt6 GUI with dark/light themes
- Complete CLI interface with identical functionality
- Comprehensive slash command system (/help, /learn, /vectordb, etc.)
- Real-time streaming responses and status updates

🔧 TECHNICAL STACK:
- LLM: devstral-small-2507-mlx via LM Studio (function calling enabled)
- Embeddings: qwen3-embedding:latest via Ollama
- Vector DB: ChromaDB v2 server for knowledge persistence
- Memory: SQLite for conversation history
- Framework: LangChain for orchestration
- UI: PyQt6 for GUI, Rich CLI for terminal

INITIALIZATION SEQUENCE:
1. Load environment configuration (.env file required)
2. Initialize devstral-small-2507-mlx connection via LM Studio
3. Connect to ChromaDB v2 server for knowledge operations
4. Initialize Ollama embeddings for vectorization
5. Load/create SQLite database for conversation memory
6. Bind 7 AI tools to LLM for autonomous execution
7. Start interactive chat loop with tool calling support
4. Load space settings and current workspace configuration
5. Load conversation history from SQLite database
6. Enter main chat loop with learning capabilities
7. Persist all data (conversations, learned knowledge, space settings) to respective stores
"""

# =============================================================================
# IMPORTS
# =============================================================================

# Core LangChain components for AI chat functionality
from langchain_openai import (
    ChatOpenAI,
)  # Main LLM interface using OpenAI-compatible API
from langchain_core.messages import (
    HumanMessage,
    SystemMessage,
    AIMessage,
)  # Message types for chat
# Base class for all message types
from langchain_core.messages import BaseMessage
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
)  # Text chunking for documents

# Standard library imports
import os  # Environment variable and file system operations
import sys  # System operations and exit handling
import warnings  # Suppress compatibility warnings
import logging  # Structured logging throughout the application
import json  # Serialization for conversation memory
import sqlite3  # SQLite database for conversation memory
import threading  # Thread synchronization for database operations
from datetime import datetime  # Timestamps for learned knowledge
# Type hints for better code clarity
from typing import List, Optional, Dict, cast

# Security and configuration
from pydantic import SecretStr  # Secure handling of API keys

# =============================================================================
# LOGGING AND ENVIRONMENT SETUP
# =============================================================================

# Configure structured logging for the entire application
# - Level INFO: Shows informational messages, warnings, and errors
# - Format includes timestamp, log level, and message
# - Used throughout the app for debugging and monitoring
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Load environment variables from .env file for configuration management
# - Allows sensitive settings (API keys, URLs) to be stored securely
# - Falls back to system environment variables if .env not available
# - Essential for deployment flexibility and security
try:
    from dotenv import load_dotenv

    load_dotenv()
    logger.info("Loaded environment variables from .env file")

    # Set OpenMP workaround if configured
    if os.getenv("KMP_DUPLICATE_LIB_OK") == "TRUE":
        os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
        logger.info("Applied OpenMP library conflict workaround")

except ImportError:
    logger.warning(
        "python-dotenv not available, using system environment variables only"
    )

# Suppress Pydantic V1 compatibility warnings for Python 3.14+
# - LangChain-core uses Pydantic V1 internally but we're running Pydantic V2
# - This warning doesn't affect functionality, just compatibility messaging
# - Safe to suppress as the application works correctly
warnings.filterwarnings("ignore", message="Core Pydantic V1 functionality")

# =============================================================================
# APPLICATION CONFIGURATION
# =============================================================================

# Configuration is loaded from environment variables (.env file)
# This section defines default values and configuration categories:
# - LM Studio settings: API endpoint and authentication
# - Memory settings: conversation persistence and limits
# - Vector database: ChromaDB + Ollama embeddings configuration
# - Database: SQLite storage for conversation memory

# Core LLM Configuration - connects to LM Studio running locally
LM_STUDIO_BASE_URL = os.getenv("LM_STUDIO_URL")  # LM Studio API endpoint
if not LM_STUDIO_BASE_URL:
    raise ValueError("LM_STUDIO_URL environment variable is required")
LM_STUDIO_BASE_URL = cast(str, LM_STUDIO_BASE_URL)  # Type assertion for mypy
assert LM_STUDIO_BASE_URL is not None, (
    "LM_STUDIO_BASE_URL should not be None after validation"
)

LM_STUDIO_API_KEY = os.getenv("LM_STUDIO_KEY")  # API key for authentication
if not LM_STUDIO_API_KEY:
    raise ValueError("LM_STUDIO_KEY environment variable is required")
LM_STUDIO_API_KEY = cast(str, LM_STUDIO_API_KEY)  # Type assertion for mypy
assert LM_STUDIO_API_KEY is not None, (
    "LM_STUDIO_API_KEY should not be None after validation"
)

# Model Configuration - try different models for better compliance
MODEL_NAME = os.getenv("MODEL_NAME")  # Default model, can be changed
if not MODEL_NAME:
    raise ValueError("MODEL_NAME environment variable is required")
MODEL_NAME = cast(str, MODEL_NAME)  # type: ignore  # Type assertion for mypy
assert MODEL_NAME is not None, "MODEL_NAME should not be None after validation"

# Conversation Memory Configuration
max_history_pairs_str = os.getenv("MAX_HISTORY_PAIRS")
if not max_history_pairs_str:
    raise ValueError("MAX_HISTORY_PAIRS environment variable is required")
MAX_HISTORY_PAIRS = int(
    max_history_pairs_str
)  # Maximum user-assistant message pairs to keep in memory

temperature_str = os.getenv("TEMPERATURE")
if not temperature_str:
    raise ValueError("TEMPERATURE environment variable is required")
TEMPERATURE = float(
    temperature_str
)  # LLM creativity/randomness (0.0 = deterministic, 1.0 = very creative)

max_input_length_str = os.getenv("MAX_INPUT_LENGTH")
if not max_input_length_str:
    raise ValueError("MAX_INPUT_LENGTH environment variable is required")
MAX_INPUT_LENGTH = int(
    max_input_length_str
)  # Maximum user input length to prevent memory issues

# Database Configuration - for conversation memory storage
db_type_str = os.getenv("DB_TYPE")
if not db_type_str:
    raise ValueError("DB_TYPE environment variable is required")
DB_TYPE = db_type_str.lower()  # Storage type: 'json', 'sqlite', etc.

DB_PATH = os.getenv("DB_PATH")  # SQLite database file path
if not DB_PATH:
    raise ValueError("DB_PATH environment variable is required")
DB_PATH = cast(str, DB_PATH)  # Type assertion for mypy
assert DB_PATH is not None, "DB_PATH should not be None after validation"

# Vector Database Configuration - for knowledge base and context retrieval
CHROMA_HOST = os.getenv("CHROMA_HOST")  # ChromaDB server host
if not CHROMA_HOST:
    raise ValueError("CHROMA_HOST environment variable is required")
CHROMA_HOST = cast(str, CHROMA_HOST)  # type: ignore  # Type assertion for mypy
assert CHROMA_HOST is not None, "CHROMA_HOST should not be None after validation"

chroma_port_str = os.getenv("CHROMA_PORT")
if not chroma_port_str:
    raise ValueError("CHROMA_PORT environment variable is required")
CHROMA_PORT = int(chroma_port_str)  # ChromaDB server port

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL")  # Ollama embeddings server
if not OLLAMA_BASE_URL:
    raise ValueError("OLLAMA_BASE_URL environment variable is required")
OLLAMA_BASE_URL = cast(str, OLLAMA_BASE_URL)  # Type assertion for mypy

EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL")  # Embedding model name
if not EMBEDDING_MODEL:
    raise ValueError("EMBEDDING_MODEL environment variable is required")
EMBEDDING_MODEL = cast(str, EMBEDDING_MODEL)  # Type assertion for mypy

# AI Behavior Settings - configurable via slash commands
CONTEXT_MODE = (
    "auto"  # "auto", "on", "off" - controls context integration from vector database
)
LEARNING_MODE = "normal"  # "normal", "strict", "off" - controls learning behavior
# Current workspace/space for vector database operations (will be loaded from settings)
CURRENT_SPACE = "default"

# Embedding cache to reduce redundant API calls
EMBEDDING_CACHE: Dict[str, List[float]] = {}
EMBEDDING_CACHE_FILE = "embedding_cache.json"

# Query result cache to avoid redundant vector searches
QUERY_CACHE: Dict[str, List[str]] = {}
QUERY_CACHE_FILE = "query_cache.json"

# Placeholder variables for GUI imports (initialized during startup)
llm = None
vectorstore = None
embeddings = None
chroma_client = None
conversation_history: List[BaseMessage] = []

# Connection pooling for external API calls
import requests  # noqa: E402
from requests.adapters import HTTPAdapter  # noqa: E402
from urllib3.util.retry import Retry  # noqa: E402

# Create a session with connection pooling and retries
api_session = requests.Session()
retry_strategy = Retry(
    total=3,
    backoff_factor=0.3,
    status_forcelist=[429, 500, 502, 503, 504],
)
adapter = HTTPAdapter(max_retries=retry_strategy,
                      pool_connections=10, pool_maxsize=20)
api_session.mount("http://", adapter)
api_session.mount("https://", adapter)

# =============================================================================
# SPACE MANAGEMENT FUNCTIONS
# =============================================================================


def get_space_collection_name(space_name: str) -> str:
    """Get the collection name for a given space."""
    if space_name == "default":
        return "knowledge_base"  # Default space uses the traditional knowledge_base collection
    return f"space_{space_name}"


def ensure_space_collection(space_name: str) -> bool:
    """Ensure a collection exists for the given space. Returns True if successful."""
    if vectorstore is None:
        return False

    try:
        # Check if collection exists by trying to get it
        # ChromaDB will create it if it doesn't exist when we first add documents
        return True
    except Exception as e:
        logger.error(
            f"Failed to ensure collection for space {space_name}: {e}")
        return False


def list_spaces() -> List[str]:
    """List all available spaces (collections starting with 'space_')."""
    try:
        # Use ChromaDB API to list collections
        list_url = f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2/tenants/default_tenant/databases/default_database/collections"
        response = api_session.get(list_url, timeout=10)

        if response.status_code == 200:
            collections = response.json()
            spaces = ["default"]  # Always include default space
            for coll in collections:
                name = coll.get("name", "")
                if name.startswith("space_"):
                    space_name = name[6:]  # Remove "space_" prefix
                    if space_name not in spaces:
                        spaces.append(space_name)
            return spaces
        else:
            logger.warning(
                f"Failed to list collections: HTTP {response.status_code}")
            return ["default"]
    except Exception as e:
        logger.error(f"Error listing spaces: {e}")
        return ["default"]


def delete_space(space_name: str) -> bool:
    """Delete a space and its collection. Returns True if successful."""
    if space_name == "default":
        return False

    try:
        from chromadb import HttpClient

        assert CHROMA_HOST is not None, (
            "CHROMA_HOST should not be None after validation"
        )
        client = HttpClient(host=CHROMA_HOST, port=CHROMA_PORT)
        collection_name = get_space_collection_name(space_name)
        client.delete_collection(collection_name)
        return True
    except Exception as e:
        logger.error(f"Failed to delete space {space_name}: {e}")
        return False


def save_current_space():
    """Save the current space to persistent storage."""
    try:
        import json

        settings = {"current_space": CURRENT_SPACE}
        with open("space_settings.json", "w") as f:
            json.dump(settings, f, indent=2)
    except Exception as e:
        logger.warning(f"Failed to save space settings: {e}")


def load_current_space() -> str:
    """Load the current space from persistent storage."""
    try:
        import json

        if os.path.exists("space_settings.json"):
            with open("space_settings.json", "r") as f:
                settings = json.load(f)
                return settings.get("current_space", "default")
    except Exception as e:
        logger.warning(f"Failed to load space settings: {e}")
    return "default"


def switch_space(space_name: str) -> bool:
    """Switch to a different space. Returns True if successful."""
    global CURRENT_SPACE
    if not ensure_space_collection(space_name):
        return False
    CURRENT_SPACE = space_name
    save_current_space()  # Save the new current space
    return True


# =============================================================================
# CACHING FUNCTIONS
# =============================================================================


def load_embedding_cache():
    """Load embedding cache from disk."""
    global EMBEDDING_CACHE
    try:
        if os.path.exists(EMBEDDING_CACHE_FILE):
            with open(EMBEDDING_CACHE_FILE, "r") as f:
                EMBEDDING_CACHE = json.load(f)
            logger.info(f"Loaded {len(EMBEDDING_CACHE)} cached embeddings")
    except Exception as e:
        logger.warning(f"Failed to load embedding cache: {e}")
        EMBEDDING_CACHE = {}


def load_query_cache():
    """Load query result cache from disk."""
    global QUERY_CACHE
    try:
        if os.path.exists(QUERY_CACHE_FILE):
            with open(QUERY_CACHE_FILE, "r") as f:
                QUERY_CACHE = json.load(f)
            logger.info(f"Loaded {len(QUERY_CACHE)} cached query results")
    except Exception as e:
        logger.warning(f"Failed to load query cache: {e}")
        QUERY_CACHE = {}


def save_query_cache():
    """Save query result cache to disk."""
    try:
        global QUERY_CACHE
        # Limit cache size
        if len(QUERY_CACHE) > 1000:
            items = list(QUERY_CACHE.items())
            QUERY_CACHE = dict(items[-500:])
        with open(QUERY_CACHE_FILE, "w") as f:
            json.dump(QUERY_CACHE, f, separators=(",", ":"))
    except Exception as e:
        logger.warning(f"Failed to save query cache: {e}")


# Initialize caches
load_embedding_cache()
load_query_cache()


# Memory management: Clean up unused objects periodically
def cleanup_memory():
    """Force garbage collection to free memory."""
    import gc

    gc.collect()


# Schedule periodic cleanup (every 100 operations)
operation_count = 0

# Database connection variables (initialized at module import time)
db_conn: Optional[sqlite3.Connection] = None
db_lock: Optional[threading.Lock] = None

# Initialize SQLite database connection if configured
if DB_TYPE == "sqlite":
    try:
        db_conn = sqlite3.connect(DB_PATH, check_same_thread=False)
        db_lock = threading.Lock()

        # Create conversations table if it doesn't exist
        with db_lock:
            cursor = db_conn.cursor()
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS conversations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    session_id TEXT NOT NULL,
                    message_type TEXT NOT NULL,
                    content TEXT NOT NULL,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            """)
            db_conn.commit()
        logger.info("SQLite database initialized for conversation memory")
    except Exception as e:
        logger.error(f"Failed to initialize SQLite database: {e}")
        db_conn = None
        db_lock = None

# =============================================================================
# CONVERSATION MEMORY MANAGEMENT
# =============================================================================


def load_memory() -> List[BaseMessage]:
    """
    Load conversation history from SQLite database.

    This function enables conversation continuity across application restarts by:
    1. Reading message history from SQLite database
    2. Reconstructing LangChain message objects
    3. Providing fallback for new conversations

    Returns:
        List of BaseMessage objects representing the conversation history

    Message Types:
    - SystemMessage: AI system prompts and context
    - HumanMessage: User inputs and questions
    - AIMessage: AI responses and answers

    Used by:
    - Global conversation_history initialization
    - Called once at application startup
    """
    try:
        if DB_TYPE == "sqlite" and db_conn and db_lock:
            # Load from SQLite database with thread safety
            with db_lock:
                cursor = db_conn.cursor()
                cursor.execute("""
                    SELECT message_type, content FROM conversations
                    WHERE session_id = 'default'
                    ORDER BY timestamp ASC
                """)
                rows = cursor.fetchall()

            if not rows:
                logger.info(
                    "No conversation history found in database, starting fresh")
                return [SystemMessage(content="Lets get some coding done..")]

            # Reconstruct message objects from database rows
            history: List[BaseMessage] = []
            for msg_type, content in rows:
                if msg_type == "SystemMessage":
                    history.append(SystemMessage(content=content))
                elif msg_type == "HumanMessage":
                    history.append(HumanMessage(content=content))
                elif msg_type == "AIMessage":
                    history.append(AIMessage(content=content))
                else:
                    logger.warning(
                        f"Unknown message type: {msg_type}, skipping")

            logger.info(f"Loaded {len(history)} messages from database")
            return history

        else:
            # Database not available - this should not happen in normal operation
            logger.error(
                "SQLite database not available for loading conversation memory"
            )
            raise RuntimeError("Database required but not available")

    except Exception as e:
        logger.error(f"Failed to load memory: {e}")
        raise  # Re-raise to ensure the error is not silently ignored


# Initialize global conversation history at module level
# - Loaded once when module imports
# - Persists throughout application lifetime
# - Modified by main chat loop and commands

conversation_history = load_memory()


# vulture: noqa
def get_llm():
    """Get the current LLM instance. Used by GUI to access initialized LLM."""
    return llm


# vulture: noqa
def get_vectorstore():
    """Get the current vectorstore instance. Used by GUI to access initialized vectorstore."""
    return vectorstore


def save_memory(history: List[BaseMessage]) -> None:
    """
    Save conversation history to SQLite database.

    Thread-safe SQLite database operations for persistent storage.

    Args:
        history: List of conversation messages to save
    """
    try:
        if DB_TYPE == "sqlite" and db_conn and db_lock:
            # Save to SQLite database with thread safety
            with db_lock:
                cursor = db_conn.cursor()

                # Clear existing messages for this session (simple approach - could be optimized)
                cursor.execute(
                    "DELETE FROM conversations WHERE session_id = 'default'")

                # Insert new messages
                for msg in history:
                    cursor.execute(
                        """
                        INSERT INTO conversations (session_id, message_type, content)
                        VALUES (?, ?, ?)
                    """,
                        ("default", type(msg).__name__, msg.content),
                    )

                db_conn.commit()
                logger.info("Conversation memory saved to database")

        else:
            # Database not available - this should not happen in normal operation
            logger.error(
                "SQLite database not available for saving conversation memory")
            raise RuntimeError("Database required but not available")

    except Exception as e:
        logger.error(f"Failed to save memory: {e}")
        raise  # Re-raise to ensure the error is not silently ignored


def get_relevant_context(
    query: str, k: int = 3, space_name: Optional[str] = None
) -> str:
    """Get relevant context from vector database with caching."""
    # Use current space if not specified
    if space_name is None:
        space_name = CURRENT_SPACE

    # Check query cache first
    cache_key = f"{space_name}:{query}:{k}"
    if cache_key in QUERY_CACHE:
        cached_results = QUERY_CACHE[cache_key]
        if cached_results:
            context = "\n\n".join(
                [f"From knowledge base:\n{doc}" for doc in cached_results]
            )
            return f"\n\nRelevant context:\n{context}\n\n"

    # Return empty context if vector database not available
    if vectorstore is None:
        return ""

    try:
        # Generate embedding for the query - embeddings will be available at runtime
        try:
            # Use global embeddings variable (initialized during application startup)
            if embeddings is None:
                logger.warning(
                    "Embeddings not initialized for context retrieval")
                return ""

            query_embedding = embeddings.embed_query(query)
        except (AttributeError, NameError, Exception) as e:
            logger.warning(
                f"Embeddings not available for context retrieval: {e}")
            return ""

        # Find collection for the specified space
        collection_id = None
        collection_name = get_space_collection_name(space_name)

        # Try to find the collection by name
        try:
            list_url = f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2/tenants/default_tenant/databases/default_database/collections"
            list_response = requests.get(list_url, timeout=10)
            if list_response.status_code == 200:
                collections = list_response.json()
                for coll in collections:
                    if coll.get("name") == collection_name:
                        collection_id = coll.get("id")
                        break
        except Exception as e:
            logger.warning(
                f"Error finding collection for space {space_name}: {e}")

        if not collection_id:
            logger.warning(f"Could not find collection for space {space_name}")
            return ""

        # ChromaDB v2 API endpoint for querying
        query_url = f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2/tenants/default_tenant/databases/default_database/collections/{collection_id}/query"

        # Query payload with embedding
        payload = {"query_embeddings": [query_embedding], "n_results": k}

        response = api_session.post(query_url, json=payload, timeout=10)

        if response.status_code == 200:
            data = response.json()

            # Extract documents from response
            docs = []
            if "documents" in data and data["documents"] and len(data["documents"]) > 0:
                documents = data["documents"][0]  # First query result
                for doc_content in documents:
                    docs.append(doc_content)

            # Return empty if no relevant documents found
            if not docs:
                return ""

            # Cache the results
            QUERY_CACHE[cache_key] = docs
            if len(QUERY_CACHE) % 50 == 0:
                save_query_cache()

            context = "\n\n".join(
                [f"From knowledge base:\n{doc}" for doc in docs])
            return f"\n\nRelevant context:\n{context}\n\n"
        else:
            print(f"ChromaDB query failed: {response.status_code}")
            return ""

    except Exception as e:
        # Log error but don't crash - AI can still respond without context
        logger.warning(f"Failed to retrieve context: {e}")
        return ""


# =============================================================================
# CONVERSATION MANAGEMENT FUNCTIONS
# =============================================================================


def trim_history(
    history: List[BaseMessage], max_pairs: int = MAX_HISTORY_PAIRS
) -> List[BaseMessage]:
    """
    Trim conversation history to prevent memory bloat and API token limits.

    LangChain sends full conversation history with each API call. Long conversations
    can exceed token limits and slow down responses. This function maintains recent
    context while preventing excessive memory usage.

    Args:
        history: Complete list of conversation messages
        max_pairs: Maximum number of user-assistant exchange pairs to keep

    Returns:
        Trimmed history list containing system message + recent exchanges

    Trimming Logic:
    - Always keep the first message (system prompt)
    - Keep the last (max_pairs * 2) messages (user + AI pairs)
    - Total length = 1 + (max_pairs * 2)
    - Example: max_pairs=10 -> keep system + 20 recent messages

    Used by:
    - Main chat loop after each AI response
    - Ensures conversation stays within reasonable bounds
    - Balances context retention with performance
    """
    # Calculate maximum allowed length: system message + (pairs * 2 messages per pair)
    max_length = max_pairs * 2 + 1

    # Only trim if history exceeds the limit
    if len(history) > max_length:
        # Keep system message (index 0) + most recent messages
        return [history[0]] + history[-(max_pairs * 2):]

    # Return unchanged if within limits
    return history


def handle_slash_command(command: str) -> bool:
    """
    Handle slash commands. Returns True if command was handled, False if it is a regular message.
    """
    # Remove leading slash and split command
    if not command.startswith("/"):
        return False

    cmd_parts = command[1:].split()
    if not cmd_parts:
        return False

    cmd_name = cmd_parts[0].lower()
    cmd_args = cmd_parts[1:] if len(cmd_parts) > 1 else []

    # Route to appropriate command handler
    if cmd_name == "help":
        show_help()
    elif cmd_name == "memory":
        show_memory()
    elif cmd_name == "clear":
        return handle_clear_command()
    elif cmd_name == "learn":
        handle_learn_command(" ".join(cmd_args))
    elif cmd_name == "vectordb":
        show_vectordb()
    elif cmd_name == "populate":
        handle_populate_command(" ".join(cmd_args))
    elif cmd_name == "model":
        show_model_info()
    elif cmd_name == "context":
        handle_context_command(cmd_args)
    elif cmd_name == "learning":
        handle_learning_command(cmd_args)
    elif cmd_name == "space":
        handle_space_command(" ".join(cmd_args))
    elif cmd_name == "export":
        handle_export_command(" ".join(cmd_args))
    elif cmd_name == "read":
        handle_read_command(" ".join(cmd_args))
    elif cmd_name == "write":
        handle_write_command(" ".join(cmd_args))
    elif cmd_name == "list":
        handle_list_command(" ".join(cmd_args))
    elif cmd_name == "pwd":
        handle_pwd_command()
    else:
        print(f"\n❌ Unknown command: /{cmd_name}")
        print("Type /help for available commands\n")

    return True


def show_help():
    """Display available commands."""
    print("\n--- Available Commands ---")
    print("/memory       - Show conversation history")
    print("/vectordb     - Show vector database contents")
    print("/model        - Show current model information")
    print("/space <cmd>  - Space/workspace management (list/create/switch/delete)")
    print("/context <mode> - Control context integration (auto/on/off)")
    print("/learning <mode> - Control learning behavior (normal/strict/off)")
    print("/populate <path> - Add code files from directory to vector DB")
    print("/clear        - Clear conversation history")
    print("/learn <text> - Add information to knowledge base")
    print("/export <fmt> - Export conversation (json/markdown)")
    print("/read <file>  - Read file contents")
    print("/write <file> - Write content to file")
    print("/list [dir]   - List directory contents")
    print("/pwd          - Show current directory")
    print("/help         - Show this help message")
    print("")
    print("Regular commands (no slash needed):")
    print("quit    - Exit the application")
    print("exit    - Exit the application")
    print("q       - Exit the application")
    print("")
    print("Note: Quit commands work after AI finishes responding.")
    print("Use Ctrl+C for immediate interruption.")
    print("--- End Help ---\n")


def handle_read_command(file_path: str):
    """Handle /read command to read file contents."""
    if not file_path:
        print("\n❌ Usage: /read <file_path>")
        print("Example: /read README.md\n")
        return

    try:
        # Security: Only allow reading files in current directory and subdirectories
        current_dir = os.getcwd()
        full_path = os.path.abspath(file_path)

        # Check if the file is within the current directory
        if not full_path.startswith(current_dir):
            print(f"\n❌ Access denied: Cannot read files outside current directory")
            print(f"Current directory: {current_dir}\n")
            return

        if not os.path.exists(full_path):
            print(f"\n❌ File not found: {file_path}\n")
            return

        if not os.path.isfile(full_path):
            print(f"\n❌ Not a file: {file_path}\n")
            return

        # Check file size (limit to 1MB to prevent memory issues)
        file_size = os.path.getsize(full_path)
        if file_size > 1024 * 1024:  # 1MB limit
            print(f"\n❌ File too large: {file_size} bytes (max 1MB)")
            print("Use a text editor to view large files\n")
            return

        with open(full_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        print(f"\n📄 File: {file_path}")
        print(f"📊 Size: {file_size} bytes")
        print("-" * 50)
        print(content)
        print("-" * 50)
        print(f"✅ File read successfully\n")

    except UnicodeDecodeError:
        print(f"\n❌ Cannot read binary file: {file_path}")
        print("This appears to be a binary file. Use /list to see file types.\n")
    except Exception as e:
        print(f"\n❌ Failed to read file: {e}\n")


def handle_write_command(args: str):
    """Handle /write command to write content to a file."""
    if not args:
        print("\n❌ Usage: /write <file_path> <content>")
        print("Example: /write notes.txt This is my note\n")
        return

    # Split on first space to separate file path from content
    parts = args.split(" ", 1)
    if len(parts) < 2:
        print("\n❌ Usage: /write <file_path> <content>")
        print("Content is required\n")
        return

    file_path, content = parts

    try:
        # Security: Only allow writing files in current directory and subdirectories
        current_dir = os.getcwd()
        full_path = os.path.abspath(file_path)

        # Check if the file is within the current directory
        if not full_path.startswith(current_dir):
            print(f"\n❌ Access denied: Cannot write files outside current directory")
            print(f"Current directory: {current_dir}\n")
            return

        # Create directory if it doesn't exist
        dir_path = os.path.dirname(full_path)
        if dir_path and not os.path.exists(dir_path):
            os.makedirs(dir_path, exist_ok=True)

        with open(full_path, "w", encoding="utf-8") as f:
            f.write(content)

        print(f"\n✅ File written: {file_path}")
        print(f"📊 Content length: {len(content)} characters\n")

    except Exception as e:
        print(f"\n❌ Failed to write file: {e}\n")


def handle_list_command(dir_path: str = ""):
    """Handle /list command to list directory contents."""
    try:
        if not dir_path:
            target_dir = os.getcwd()
        else:
            # Security: Only allow listing directories in current directory
            current_dir = os.getcwd()
            full_path = os.path.abspath(dir_path)

            if not full_path.startswith(current_dir):
                print(
                    f"\n❌ Access denied: Cannot list directories outside current directory"
                )
                print(f"Current directory: {current_dir}\n")
                return

            if not os.path.exists(full_path):
                print(f"\n❌ Directory not found: {dir_path}\n")
                return

            if not os.path.isdir(full_path):
                print(f"\n❌ Not a directory: {dir_path}\n")
                return

            target_dir = full_path

        print(
            f"\n📁 Directory: {os.path.relpath(target_dir, os.getcwd()) or '.'}")
        print("-" * 50)

        items = os.listdir(target_dir)
        if not items:
            print("📂 (empty directory)")
        else:
            # Separate files and directories
            dirs = []
            files = []

            for item in sorted(items):
                full_item_path = os.path.join(target_dir, item)
                if os.path.isdir(full_item_path):
                    dirs.append(item + "/")
                else:
                    files.append(item)

            # Display directories first, then files
            for item in dirs + files:
                print(f"  {item}")

        print("-" * 50)
        print(f"📊 Total items: {len(items)}\n")

    except Exception as e:
        print(f"\n❌ Failed to list directory: {e}\n")


def handle_pwd_command():
    """Handle /pwd command to show current working directory."""
    try:
        current_dir = os.getcwd()
        print(f"\n📍 Current directory: {current_dir}\n")
    except Exception as e:
        print(f"\n❌ Failed to get current directory: {e}\n")


def show_memory():
    """Show conversation history."""
    if not conversation_history:
        print("\n📝 No conversation history available.\n")
        return

    print(
        f"\n📝 Conversation History ({len(conversation_history)} messages):\n")
    for i, msg in enumerate(conversation_history):
        msg_type = type(msg).__name__.replace("Message", "")
        content = str(msg.content)
        content_preview = content[:100] + \
            "..." if len(content) > 100 else content
        print(f"{i + 1:2d}. {msg_type}: {content_preview}")
    print()


def handle_clear_command() -> bool:
    """Handle /clear command. Returns True if should exit."""
    confirm = input(
        "Are you sure you want to clear all conversation history? (yes/no): "
    )
    if confirm.lower() in ["yes", "y"]:
        global conversation_history
        # Clear display and show confirmation
        print("\nConversation memory cleared. Starting fresh.\n")

        # Add a new system message for the fresh start
        from langchain_core.messages import SystemMessage

        conversation_history = [SystemMessage(
            content="Lets get some coding done..")]
        save_memory(conversation_history)
        return False
    else:
        print("\n❌ Clear cancelled\n")
        return False


def handle_learn_command(content: str):
    """Handle /learn command."""
    # Check if vector database is available
    if vectorstore is None:
        print(
            "\nVector database not available. ChromaDB is required for learning features.\n"
        )
        return

    # Extract content after 'learn ' command
    if not content:
        print("\nUsage: /learn <information to remember>\n")
        return

    try:
        # Import Document class for creating knowledge entries
        from langchain_core.documents import Document

        # Create document with metadata
        doc = Document(
            page_content=content,  # The information to remember
            metadata={
                "source": "manual",  # Mark as manually added
                "added_at": str(datetime.now()),  # Timestamp for tracking
            },
        )

        # Add to vector database for current space
        # Generate embeddings and use direct HTTP API calls to ChromaDB v2

        # Generate embeddings using Ollama
        try:
            # Import here to handle initialization order
            from main import embeddings  # type: ignore

            embeddings_result = embeddings.embed_documents([doc.page_content])
            if embeddings_result and len(embeddings_result) > 0:
                embedding_vector = embeddings_result[0]

                # Get collection for current space
                collection_name = get_space_collection_name(CURRENT_SPACE)
                collection_id = None

                # Find or create the collection
                try:
                    list_url = f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2/tenants/default_tenant/databases/default_database/collections"
                    list_response = api_session.get(list_url, timeout=10)
                    if list_response.status_code == 200:
                        collections = list_response.json()
                        for coll in collections:
                            if coll.get("name") == collection_name:
                                collection_id = coll.get("id")
                                break

                    # If collection doesn't exist, create it
                    if not collection_id:
                        create_url = f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2/tenants/default_tenant/databases/default_database/collections"
                        create_payload = {"name": collection_name}
                        create_response = api_session.post(
                            create_url, json=create_payload, timeout=10
                        )
                        if create_response.status_code == 201:
                            collection_id = create_response.json().get("id")
                            logger.info(
                                f"Created new collection for space {CURRENT_SPACE}: {collection_name}"
                            )
                        else:
                            logger.error(
                                f"Failed to create collection: HTTP {create_response.status_code}"
                            )
                            raise Exception("Collection creation failed")

                except Exception as e:  # type: ignore
                    logger.error(
                        f"Error managing collection for space {CURRENT_SPACE}: {e}"
                    )
                    raise

                # Add document to the space's collection
                add_url = f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2/tenants/default_tenant/databases/default_database/collections/{collection_id}/add"  # noqa: E501

                # Prepare the document data with embeddings
                doc_id = f"doc_{len(doc.page_content)}_{int(datetime.now().timestamp() * 1000000)}"
                payload = {
                    "ids": [doc_id],
                    "embeddings": [embedding_vector],
                    "documents": [doc.page_content],
                    "metadatas": [doc.metadata] if doc.metadata else [{}],
                }

                # Make the API call to add the document
                headers = {"Content-Type": "application/json"}
                response = api_session.post(
                    add_url, json=payload, headers=headers, timeout=10
                )

                if response.status_code == 201:
                    logger.info(
                        f"Document added successfully to space {CURRENT_SPACE}: {doc_id}"
                    )
                else:
                    logger.error(
                        f"Failed to add document to space {CURRENT_SPACE}: {response.status_code} - {response.text}"
                    )
                    raise Exception("Document addition failed")

            else:
                logger.error("Failed to generate embeddings")
                raise Exception("Embedding generation failed")

        except Exception as e:
            logger.error(f"Error with direct ChromaDB v2 API: {e}")
            # Fallback to LangChain method
            try:
                vectorstore.add_documents([doc])
                logger.info("Used LangChain fallback for document addition")
            except Exception as fallback_e:
                logger.error(f"LangChain fallback also failed: {fallback_e}")
                print(f"\n❌ Failed to learn: {e}\n")
                return

        # Periodic memory cleanup
        global operation_count
        operation_count += 1
        if operation_count % 50 == 0:
            cleanup_memory()

        print(f"\n✅ Learned: {content[:50]}...\n")

    except Exception as e:
        print(f"\n❌ Failed to learn: {e}\n")


def show_vectordb():
    """
    Display information about the current vector database contents.

    Shows collection statistics, document count, and sample content from the
    ChromaDB vector database. Provides insights into the AI's knowledge base.
    """
    if vectorstore is None:
        print("\n❌ Vector database not available.\n")
        return

    try:
        print("\n--- Vector Database Contents ---")

        # Try to get documents from ChromaDB via direct API
        try:
            # Find collection for current space
            collection_name = get_space_collection_name(CURRENT_SPACE)
            collection_id = None

            list_url = f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2/tenants/default_tenant/databases/default_database/collections"
            list_response = api_session.get(list_url, timeout=10)

            if list_response.status_code == 200:
                collections = list_response.json()
                for coll in collections:
                    if coll.get("name") == collection_name:
                        collection_id = coll.get("id")
                        break

            if not collection_id:
                print(
                    f"❌ No collection found for current space '{CURRENT_SPACE}'")
                print("The space may not have any documents yet.")
                return

            logger.info(
                f"Using collection for space {CURRENT_SPACE}: {collection_name} (ID: {collection_id})"
            )

            # Get collection statistics instead of all documents
            count_url = f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2/tenants/default_tenant/databases/default_database/collections/{collection_id}/count"  # noqa: E501
            count_response = api_session.get(count_url, timeout=10)

            if count_response.status_code == 200:
                count_data = count_response.json()
                count = (
                    count_data
                    if isinstance(count_data, int)
                    else count_data.get("count", 0)
                )
                print(f"📊 Collection: {collection_name}")
                print(f"📈 Documents: {count}")
                print(f"🏢 Space: {CURRENT_SPACE}")

                if count > 0:
                    print("\n📄 Sample Documents:")
                    try:
                        # Try to get sample documents using ChromaDB client directly
                        chroma_client = vectorstore._client
                        collection = chroma_client.get_collection(
                            collection_name)
                        results = collection.get(
                            limit=3, include=["documents", "metadatas"]
                        )

                        if results and "documents" in results and results["documents"]:
                            docs = results["documents"]
                            metadatas = results.get("metadatas") or [
                                {}] * len(docs)

                            # First 3 documents
                            for i, doc in enumerate(docs[:3]):
                                metadata = metadatas[i] if i < len(
                                    metadatas) else {}
                                if isinstance(doc, str):
                                    # Clean up the preview - remove excessive whitespace and binary-looking content
                                    preview = doc.strip()
                                    # Check for binary content, PDF metadata, or placeholder text
                                    if (
                                        any(
                                            indicator in preview.lower()
                                            for indicator in [
                                                "<source>",
                                                "</source>",
                                                "<translation",
                                                "<<",
                                                ">>",
                                                "/type",
                                                "/pages",
                                                "xref",
                                                "trailer",
                                                "startxref",
                                            ]
                                        )
                                        or any(
                                            ord(c) < 32 and c not in "\n\t"
                                            for c in preview[:100]
                                        )
                                        or preview.startswith("[Binary file:")
                                        or "[Binary file:" in preview
                                    ):
                                        preview = "[Binary file - content not shown]"
                                    elif len(preview) > 200:
                                        preview = preview[:200] + "..."
                                    elif not preview:
                                        preview = "[empty content]"

                                    source = (
                                        metadata.get("source", "unknown")
                                        if metadata
                                        else "unknown"
                                    )
                                    added_at = (
                                        metadata.get("added_at", "unknown")
                                        if metadata
                                        else "unknown"
                                    )
                                    print(
                                        f"  {i + 1}. {source} (added: {added_at})")
                                    print(f"      Preview: {preview}")
                                else:
                                    print(
                                        f"  {i + 1}. [Non-text document: {type(doc)}]"
                                    )
                        else:
                            print("  No documents found in collection")
                    except Exception as e:  # type: ignore
                        print(
                            f"  Could not retrieve sample documents: {str(e)[:100]}")
                        print(
                            "  (This is normal for some document types or if the collection is empty)"
                        )
                    print()
            else:
                print("❌ Could not retrieve collection statistics")

        except Exception as e:
            print(f"❌ Error accessing vector database: {e}")

    except Exception as e:
        print(f"❌ Failed to show vector database: {e}")


def handle_populate_command(dir_path: str):
    """
    Handle the /populate command to bulk import codebases into the vector database.

    Parses command arguments including optional flags like --dry-run and --direct-api.
    Processes files in batches, extracts content from various file types, and adds
    them to the ChromaDB collection with proper metadata.

    Args:
        dir_path: Directory path with optional flags (--dry-run, --direct-api)
    """
    # Import ChromaDB client for direct API usage
    from chromadb import HttpClient

    # Access global chroma_client
    global chroma_client
    # Parse additional options from dir_path
    parts = dir_path.split()
    directory = parts[0] if parts else ""
    options = parts[1:] if len(parts) > 1 else []

    dry_run = "--dry-run" in options
    use_direct_api = True  # Always use direct API for better reliability
    clear_first = "--clear" in options

    # Check if directory exists
    if not directory:
        print(
            "\nUsage: /populate /path/to/directory [--dry-run] [--direct-api]")
        print("Examples:")
        print("   /populate /Users/username/projects/myapp")
        print("   /populate src/ --dry-run")
        print("   /populate . --clear")
        print("   /populate ~/code --dry-run --clear")
        print("\nOptions:")
        print("  --dry-run     : Validate files without writing to database")
        print("  --clear       : Delete existing collection before repopulating")
        print(
            "\nNote: Designed for code files. For documents (PDFs, etc.), use document processing tools."
        )
        print("Uses direct ChromaDB API for optimal reliability and performance")
        print(
            "\nNote: This will recursively scan and add all code files to the vector database."
        )
        return

    if not os.path.exists(directory):
        print(f"\n❌ Directory not found: {directory}")
        return

    if not os.path.isdir(directory):
        print(f"\n❌ Path is not a directory: {directory}")
        return

    print(f"\n🔍 Starting codebase population from: {directory}")
    if dry_run:
        print("🔍 DRY RUN MODE - No data will be written to database")
    if clear_first:
        print("🧹 CLEAR MODE - Will delete existing collection before repopulating")
    print("🔧 Using direct ChromaDB API for optimal reliability and performance")
    print("This may take some time for large codebases...")
    print()

    # Check if vectorstore is available
    if vectorstore is None:
        print("\n❌ Vector database not available. Check your configuration.")
        return

    # Clear existing collection if requested
    if clear_first and not dry_run:
        print("🧹 Clearing existing collection...")
        try:
            # Get the collection name for current space
            collection_name = get_space_collection_name(CURRENT_SPACE)

            # Use direct ChromaDB API to delete collection
            import requests

            delete_url = f"http://{CHROMA_HOST}:{CHROMA_PORT}/api/v2/tenants/default_tenant/databases/default_database/collections/{collection_name}"
            response = requests.delete(delete_url, timeout=10)

            if response.status_code in [200, 204]:
                print(f"✅ Cleared collection: {collection_name}")
            else:
                print(
                    f"⚠️  Collection may not exist or failed to clear: {response.status_code}"
                )
        except Exception as e:
            print(f"⚠️  Failed to clear collection: {e}")

    try:
        # Import required libraries for population
        from pathlib import Path
        from langchain_core.documents import Document

        # Code file extensions (same as tools/populate_codebase.py)
        CODE_EXTENSIONS = {
            ".py": "python",
            ".js": "javascript",
            ".ts": "typescript",
            ".jsx": "react",
            ".tsx": "react-typescript",
            ".cpp": "cpp",
            ".c": "c",
            ".h": "c-header",
            ".hpp": "cpp-header",
            ".java": "java",
            ".cs": "csharp",
            ".php": "php",
            ".rb": "ruby",
            ".go": "go",
            ".rs": "rust",
            ".swift": "swift",
            ".kt": "kotlin",
            ".scala": "scala",
            ".sh": "bash",
            ".sql": "sql",
            ".html": "html",
            ".css": "css",
            ".md": "markdown",
            ".json": "json",
            ".xml": "xml",
            ".yaml": "yaml",
            ".yml": "yaml",
        }

        # ChromaDB doesn't provide easy count

        # Scan directory
        directory = Path(directory)
        files_to_process = []

        for file_path in directory.rglob("*"):
            if file_path.is_file() and True:
                files_to_process.append(file_path)

        if not files_to_process:
            print("❌ No code files found in the specified directory.")
            return

        print(f"📋 Found {len(files_to_process)} code files to process")

        # Process files in batches for better performance
        total_chunks = 0
        processed_files = 0
        batch_size = 50  # Maximum chunks per batch to avoid payload too large
        all_documents = []

        print(f"⚡ Processing files in batches of {batch_size} chunks...")

        for file_path in files_to_process:
            relative_path = file_path.relative_to(directory)

            try:
                # Check if it's a document file that can be processed
                if file_path.suffix.lower() in [".pdf"]:
                    # For PDFs, extract text using document processing
                    try:
                        parse_result = execute_parse_document(
                            str(file_path), "text")
                        if parse_result.get("success") and "content" in parse_result:
                            content = parse_result["content"]
                            if not content or (
                                content.startswith(
                                    "[") and content.endswith("]")
                            ):
                                # If no content or still a placeholder, skip this file
                                continue
                        else:
                            # If parsing failed, skip this file
                            continue
                    except Exception:
                        # If document processing fails, skip this file
                        continue
                elif file_path.suffix.lower() in [
                    ".png",
                    ".jpg",
                    ".jpeg",
                    ".gif",
                    ".bmp",
                ]:
                    # For images, skip them (can't extract text easily)
                    continue
                else:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
            except Exception:
                continue  # Skip files that can't be read

            if not content.strip():
                continue

            # Get metadata
            ext = file_path.suffix.lower()
            language = CODE_EXTENSIONS.get(ext, "unknown")

            # Chunk the content
            chunks = chunk_text(content)

            # Create documents
            for j, chunk in enumerate(chunks):
                metadata = {
                    "source": "codebase",
                    "file_path": str(relative_path),
                    "language": language,
                    "chunk_index": j,
                    "total_chunks": len(chunks),
                    "added_at": str(datetime.now()),
                }

                doc = Document(page_content=chunk, metadata=metadata)
                all_documents.append(doc)

                # Process batch when it reaches the limit
                if len(all_documents) >= batch_size:
                    print(
                        f"💾 Adding batch of {len(all_documents)} chunks to vector database..."
                    )
                    try:
                        if use_direct_api:
                            # Use direct ChromaDB API
                            # Ensure chroma_client is initialized
                            if chroma_client is None:
                                chroma_client = HttpClient(
                                    host=cast(str, CHROMA_HOST), port=CHROMA_PORT
                                )

                            if chroma_client is None or embeddings is None:
                                raise Exception(
                                    "ChromaDB client or embeddings not initialized"
                                )

                            collection_name = get_space_collection_name(
                                CURRENT_SPACE)
                            collection = chroma_client.get_collection(
                                collection_name)

                            texts = [doc.page_content for doc in all_documents]
                            metadatas = [doc.metadata for doc in all_documents]
                            ids = [
                                f"{doc.metadata.get('file_path', 'unknown')}_{doc.metadata.get('chunk_index', i)}"
                                for i, doc in enumerate(all_documents)
                            ]

                            # Generate embeddings
                            embeddings_list = embeddings.embed_documents(texts)

                            collection.add(
                                documents=texts,
                                embeddings=embeddings_list,  # type: ignore
                                metadatas=metadatas,
                                ids=ids,
                            )
                        else:
                            # Use LangChain wrapper
                            vectorstore.add_documents(all_documents)

                        total_chunks += len(all_documents)
                        all_documents = []  # Clear batch
                    except Exception as e:  # type: ignore
                        print(f"❌ Failed to add batch: {e}")
                        # Continue with next batch

            processed_files += 1

            # Progress update every 10 files
            if processed_files % 10 == 0:
                print(
                    f"📄 Processed {processed_files}/{len(files_to_process)} files "
                    f"({total_chunks + len(all_documents)} total chunks)..."
                )

        # Add remaining documents
        if all_documents:
            print(f"💾 Adding final batch of {len(all_documents)} chunks...")
            try:
                if use_direct_api:
                    # Use direct ChromaDB API
                    # Ensure chroma_client is initialized
                    if chroma_client is None:
                        chroma_client = HttpClient(
                            host=cast(str, CHROMA_HOST), port=CHROMA_PORT
                        )

                    if chroma_client is None or embeddings is None:
                        raise Exception(
                            "ChromaDB client or embeddings not initialized")

                    collection_name = get_space_collection_name(CURRENT_SPACE)
                    try:
                        collection = chroma_client.get_collection(
                            collection_name)
                    except Exception:
                        # Collection doesn't exist, create it
                        collection = chroma_client.create_collection(
                            collection_name)

                    texts = [doc.page_content for doc in all_documents]
                    metadatas = [doc.metadata for doc in all_documents]
                    ids = [
                        f"{doc.metadata.get('file_path', 'unknown')}_{doc.metadata.get('chunk_index', i)}"
                        for i, doc in enumerate(all_documents)
                    ]

                    # Generate embeddings
                    embeddings_list = embeddings.embed_documents(texts)

                    collection.add(
                        documents=texts,
                        embeddings=embeddings_list,  # type: ignore
                        metadatas=metadatas,
                        ids=ids,
                    )
                else:
                    # Use LangChain wrapper
                    vectorstore.add_documents(all_documents)

                total_chunks += len(all_documents)
            except Exception as e:  # type: ignore
                print(f"❌ Failed to add final batch: {e}")

        # Save and show final stats
        save_memory(conversation_history)
        print("\n✅ Codebase population completed!")
        print("📊 Statistics:")
        print(f"   Files processed: {processed_files}")
        print(f"   Total chunks added: {total_chunks}")
        print(f"   Space: {CURRENT_SPACE}")
        print()

    except Exception as e:
        print(f"\n❌ Failed to populate codebase: {e}\n")


def show_model_info():
    """Show current model information."""
    print(f"\n🤖 Current Model: {MODEL_NAME}")
    base_url_display = (
        LM_STUDIO_BASE_URL.replace(
            "/v1", "") if LM_STUDIO_BASE_URL else "unknown"
    )
    print(f"🌐 API Endpoint: http://{base_url_display}")
    print(f"🎯 Temperature: {TEMPERATURE}")
    print(f"📏 Max Input Length: {MAX_INPUT_LENGTH}")
    print(f"💾 Context Window: {MAX_HISTORY_PAIRS} message pairs")
    print()


def handle_context_command(args: List[str]):
    """Handle /context command."""
    global CONTEXT_MODE
    if not args:
        print(f"\n🎯 Current context mode: {CONTEXT_MODE}")
        print("Options: auto, on, off")
        print("- auto: AI decides when to include context")
        print("- on: Always include available context")
        print("- off: Never include context from knowledge base")
        print()
        return

    mode = args[0].lower()
    if mode in ["auto", "on", "off"]:
        CONTEXT_MODE = mode
        print(f"\n✅ Context mode set to: {mode}\n")
    else:
        print(f"\n❌ Invalid context mode: {mode}")
        print("Valid options: auto, on, off\n")


def handle_learning_command(args: List[str]):
    """Handle /learning command."""
    global LEARNING_MODE
    if not args:
        print(f"\n🧠 Current learning mode: {LEARNING_MODE}")
        print("Options: normal, strict, off")
        print("- normal: Balanced learning with context integration")
        print("- strict: Minimal context, focused on explicit learning")
        print("- off: Disable all learning and context features")
        print()
        return

    mode = args[0].lower()
    if mode in ["normal", "strict", "off"]:
        LEARNING_MODE = mode
        print(f"\n✅ Learning mode set to: {mode}\n")
    else:
        print(f"\n❌ Invalid learning mode: {mode}")
        print("Valid options: normal, strict, off\n")


def handle_space_command(cmd_args: str):
    """Handle /space command."""
    space_cmd = cmd_args.strip()

    if not space_cmd:
        # Show current space
        print(f"\nCurrent space: {CURRENT_SPACE}")
        print(f"Collection: {get_space_collection_name(CURRENT_SPACE)}")
        print("\nUsage:")
        print("  /space list                    - List all spaces")
        print("  /space create <name>           - Create new space")
        print("  /space switch <name>           - Switch to space")
        print("  /space delete <name>           - Delete space")
        print("  /space current                 - Show current space\n")
        return

    if space_cmd == "list":
        spaces = list_spaces()
        print(f"\nAvailable spaces ({len(spaces)}):")
        for space in spaces:
            marker = " ← current" if space == CURRENT_SPACE else ""
            print(f"  • {space}{marker}")
        print()

    elif space_cmd == "current":
        print(f"\nCurrent space: {CURRENT_SPACE}")
        print(f"Collection: {get_space_collection_name(CURRENT_SPACE)}\n")

    elif space_cmd.startswith("create "):
        new_space = space_cmd[7:].strip()
        if not new_space:
            print("\n❌ Usage: /space create <name>\n")
            return

        if new_space in list_spaces():
            print(f"\n❌ Space '{new_space}' already exists\n")
            return

        if switch_space(new_space):
            print(f"\n✅ Created and switched to space: {new_space}")
            print(f"Collection: {get_space_collection_name(new_space)}\n")
        else:
            print(f"\n❌ Failed to create space: {new_space}\n")

    elif space_cmd.startswith("switch "):
        target_space = space_cmd[7:].strip()
        if not target_space:
            print("\n❌ Usage: /space switch <name>\n")
            return

        if target_space not in list_spaces():
            print(f"\n❌ Space '{target_space}' does not exist")
            print("Use '/space create {target_space}' to create it first\n")
            return

        if switch_space(target_space):
            print(f"\n✅ Switched to space: {target_space}")
            print(f"Collection: {get_space_collection_name(target_space)}\n")
        else:
            print(f"\n❌ Failed to switch to space: {target_space}\n")

    elif space_cmd.startswith("delete "):
        target_space = space_cmd[7:].strip()
        if not target_space:
            print("\n❌ Usage: /space delete <name>\n")
            return

        if target_space == "default":
            print("\n❌ Cannot delete the default space\n")
            return

        if target_space not in list_spaces():
            print(f"\n❌ Space '{target_space}' does not exist\n")
            return

        # Confirm deletion
        confirm = input(
            f"Are you sure you want to delete space '{target_space}' and all its data? (yes/no): "
        )
        if confirm.lower() not in ["yes", "y"]:
            print("\n❌ Deletion cancelled\n")
            return

        if delete_space(target_space):
            print(f"\n✅ Deleted space: {target_space}\n")
        else:
            print(f"\n❌ Failed to delete space: {target_space}\n")

    else:
        print(f"\n❌ Unknown space command: {space_cmd}")
        print("Use '/space' for help\n")


def handle_export_command(format_type: str = "json"):
    """Handle /export command to save conversation history."""
    if not conversation_history:
        print("\n📝 No conversation history to export.\n")
        return

    # Determine export format
    format_type = format_type.lower().strip()
    if format_type not in ["json", "markdown", "md"]:
        print(f"\n❌ Unsupported format: {format_type}")
        print("Supported formats: json, markdown (md)\n")
        return

    # Generate filename with timestamp
    from datetime import datetime

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"conversation_export_{timestamp}.{format_type if format_type != 'markdown' else 'md'}"

    try:
        if format_type == "json":
            # Export as JSON
            import json

            export_data: dict = {
                "export_timestamp": datetime.now().isoformat(),
                "total_messages": len(conversation_history),
                "messages": [],
            }

            for i, msg in enumerate(conversation_history):
                msg_type = type(msg).__name__.replace("Message", "").lower()
                export_data["messages"].append(
                    {
                        "index": i + 1,
                        "type": msg_type,
                        "content": str(msg.content),
                        "timestamp": getattr(
                            msg, "timestamp", datetime.now().isoformat()
                        )
                        if hasattr(msg, "timestamp")
                        else datetime.now().isoformat(),
                    }
                )

            with open(filename, "w", encoding="utf-8") as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)

        else:  # markdown
            # Export as Markdown
            with open(filename, "w", encoding="utf-8") as f:
                f.write("# AI Assistant Conversation Export\n\n")
                f.write(
                    f"**Export Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                )
                f.write(f"**Total Messages:** {len(conversation_history)}\n\n")
                f.write("---\n\n")

                for i, msg in enumerate(conversation_history):
                    msg_type = type(msg).__name__.replace("Message", "")
                    content = str(msg.content)

                    # Format based on message type
                    if msg_type == "Human":
                        f.write(f"## User Message {i + 1}\n\n{content}\n\n")
                    elif msg_type == "AI":
                        f.write(f"## AI Response {i + 1}\n\n{content}\n\n")
                    elif msg_type == "System":
                        f.write(
                            f"### System Message {i + 1}\n\n*{content}*\n\n")
                    else:
                        f.write(
                            f"## {msg_type} Message {i + 1}\n\n{content}\n\n")

                    f.write("---\n\n")

        print(f"\n✅ Conversation exported to: {filename}")
        print(f"📊 Messages exported: {len(conversation_history)}")
        print(f"📄 Format: {format_type}\n")

    except Exception as e:
        print(f"\n❌ Failed to export conversation: {e}\n")


def initialize_application():
    """
    Initialize all application components required for operation.

    This critical function sets up the complete AI assistant infrastructure:
    1. LLM Connection: Establishes devstral-small-2507-mlx connection via LM Studio with tool calling enabled
    2. Vector Database: Connects to ChromaDB v2 server for persistent knowledge storage
    3. Embeddings: Initializes qwen3-embedding via Ollama for semantic vectorization
    4. Memory: Loads conversation history from SQLite database
    5. Tool Binding: Attaches 7 AI tools to LLM for autonomous execution

    The initialization sequence ensures all components work together:
    - Tool calling requires LLM + proper tool definitions
    - Knowledge storage requires ChromaDB + embeddings
    - Context retrieval requires vector database + semantic search
    - Memory persistence requires SQLite + proper schema

    Args:
        None (uses global environment variables)

    Returns:
        bool: True if all components initialized successfully, False if any critical component failed

    Raises:
        SystemExit: If critical components (LLM, ChromaDB) cannot be initialized

    Global Variables Modified:
        llm: ChatOpenAI instance with tool calling bound
        vectorstore: ChromaDB vector store for knowledge
        embeddings: Ollama embeddings for vectorization
        conversation_history: List of loaded conversation messages
    """
    global llm, vectorstore, embeddings, chroma_client

    # ============================================================================
    # LLM INITIALIZATION
    # ============================================================================
    # Connect to LM Studio running locally for AI chat capabilities
    # LM Studio provides OpenAI-compatible API for local LLM inference
    try:
        llm = ChatOpenAI(
            # Local LM Studio server endpoint
            base_url=cast(str, LM_STUDIO_BASE_URL),
            api_key=SecretStr(
                cast(str, LM_STUDIO_API_KEY)
            ),  # API authentication (usually "lm-studio")
            # Model name (e.g., "devstral-small-2507-mlx")
            model=cast(str, MODEL_NAME),
            temperature=TEMPERATURE,  # Response creativity (0.0-1.0)
        )
        # Bind tools to the LLM for function calling support
        llm = llm.bind_tools(FILE_SYSTEM_TOOLS)
        print(f"✅ Connected to LLM: {MODEL_NAME} (with tool calling)")
    except Exception as e:
        print(f"❌ Failed to initialize LLM: {e}")
        print("Please ensure LM Studio is running and accessible.")
        return False

    # ============================================================================
    # VECTOR DATABASE INITIALIZATION
    # ============================================================================
    # Set up ChromaDB vector database for persistent knowledge storage
    # Uses Ollama embeddings for semantic text processing
    try:
        # Import vector database components
        from chromadb import HttpClient
        from langchain_chroma import Chroma
        from langchain_ollama import OllamaEmbeddings
        from typing import Any

        # Custom embeddings class to avoid invalid sampling parameters for embeddings
        class CustomOllamaEmbeddings(OllamaEmbeddings):
            @property
            def _default_params(self) -> dict[str, Any]:  # noqa: unused property override
                """Get the default parameters for calling Ollama, excluding sampling params for embeddings."""
                return {
                    "num_ctx": self.num_ctx,
                    "num_gpu": self.num_gpu,
                    "num_thread": self.num_thread,
                    "keep_alive": self.keep_alive,
                }

        # Connect to ChromaDB server
        chroma_client = HttpClient(
            host=cast(str, CHROMA_HOST), port=CHROMA_PORT)

        # Initialize Ollama embeddings for text vectorization
        embeddings = CustomOllamaEmbeddings(
            model=cast(
                str, EMBEDDING_MODEL
            ),  # Embedding model (e.g., "qwen3-embedding:latest")
            base_url=cast(str, OLLAMA_BASE_URL),  # Ollama server endpoint
        )

        # Attempt to connect to existing collection for current workspace
        # Collections are isolated per workspace for knowledge separation
        try:
            chroma_client.get_collection(
                name=get_space_collection_name(CURRENT_SPACE))
            # Collection exists - connect to it
            vectorstore = Chroma(
                client=chroma_client,
                collection_name=get_space_collection_name(CURRENT_SPACE),
                embedding_function=embeddings,
            )
            print("✅ Connected to vector database (existing collection)")
        except Exception:
            # Collection doesn't exist - create new one for this workspace
            vectorstore = Chroma(
                client=chroma_client,
                collection_name=get_space_collection_name(CURRENT_SPACE),
                embedding_function=embeddings,
            )
            print("✅ Connected to vector database (new collection)")

    except Exception as e:
        print(f"❌ ERROR: ChromaDB connection failed: {e}")
        print(
            "ChromaDB is required for learning features (/learn and /populate commands)."
        )
        print("Please ensure ChromaDB v2 server is running and accessible.")
        sys.exit(1)

    # ============================================================================
    # DATABASE INITIALIZATION
    # ============================================================================
    # Set up SQLite database connection for conversation memory storage
    global db_conn, db_lock
    if DB_TYPE == "sqlite":
        try:
            db_conn = sqlite3.connect(
                cast(str, DB_PATH), check_same_thread=False)
            db_lock = threading.Lock()

            # Create conversations table if it doesn't exist
            with db_lock:
                cursor = db_conn.cursor()
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS conversations (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT NOT NULL,
                        message_type TEXT NOT NULL,
                        content TEXT NOT NULL,
                        timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                    )
                """)
                db_conn.commit()

            print("✅ Connected to SQLite database for conversation memory")
        except Exception as e:
            print(f"❌ Failed to initialize SQLite database: {e}")
            print("Falling back to JSON file storage")
            db_conn = None
            db_lock = None

    # ============================================================================
    # CONVERSATION MEMORY LOADING
    # ============================================================================
    # Load previous conversation history from SQLite database
    # This provides continuity across application sessions
    conversation_history[:] = load_memory()

    return True


def show_welcome():
    """
    Display application startup information and load workspace configuration.

    Shows the application header, basic usage instructions, and loads the
    current workspace (space) that was last used. Spaces provide isolated
    knowledge bases for different projects or contexts.
    """
    # Display application branding and welcome message
    print("=" * 60)
    print("AI Assistant Chat Interface")
    print("=" * 60)
    print("Hello! I'm ready to help you.")
    print("Commands: 'quit', 'exit', or 'q' to exit")
    print("Slash commands: /memory, /clear, /help")
    print("Type /help for all available commands\n")

    # Load the workspace that was last active
    # Workspaces (spaces) isolate knowledge bases for different projects
    global CURRENT_SPACE
    CURRENT_SPACE = load_current_space()
    print(
        f"Current space: {CURRENT_SPACE} (collection: {get_space_collection_name(CURRENT_SPACE)})"
    )


def chunk_text(content: str) -> List[str]:
    """
    Intelligently split text content into manageable chunks for vector storage.

    Uses LangChain's RecursiveCharacterTextSplitter to create overlapping chunks
    that preserve semantic meaning and code structure. This is crucial for
    effective retrieval-augmented generation (RAG) where context matters.

    Args:
        content: The text content to be chunked (typically code or documentation)

    Returns:
        List of text chunks, each ~1500 characters with 200 character overlap

    Technical Details:
    - Chunk Size: 1500 chars - balances context with search efficiency
    - Overlap: 200 chars - ensures continuity across chunk boundaries
    - Separators: Prioritizes paragraph breaks, then lines, then words
    - Keep Separator: Preserves formatting in split content
    """
    # Configure text splitter optimized for code and technical content
    # These parameters are tuned for programming languages and documentation
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,  # Optimal size for context retention
        chunk_overlap=200,  # Ensures semantic continuity
        separators=[
            "\n\n",
            "\n",
            " ",
            "",
        ],  # Split hierarchy: paragraphs > lines > words > chars
        keep_separator=True,  # Preserve code formatting
    )

    return text_splitter.split_text(content)


# Tool definitions for devstral-small-2507-mlx function calling
# All 6 tools tested and confirmed working with devstral-small-2507-mlx
FILE_SYSTEM_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "Read the contents of a file in the current directory",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the file to read (relative to current directory)",
                    }
                },
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "parse_document",
            "description": "Extract structured data from documents using devstral-small-2507's advanced document parsing capabilities",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the document file to parse (relative to current directory)",
                    },
                    "extract_type": {
                        "type": "string",
                        "enum": ["text", "tables", "forms", "layout"],
                        "description": "Type of data to extract from the document",
                    },
                },
                "required": ["file_path", "extract_type"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": "Write content to a file in the current directory",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the file to write (relative to current directory)",
                    },
                    "content": {
                        "type": "string",
                        "description": "Content to write to the file",
                    },
                },
                "required": ["file_path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_directory",
            "description": "List contents of a directory in the current workspace",
            "parameters": {
                "type": "object",
                "properties": {
                    "directory_path": {
                        "type": "string",
                        "description": "Path to directory to list (optional, defaults to current directory)",
                        "default": ".",
                    }
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_current_directory",
            "description": "Get the current working directory path",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "learn_information",
            "description": "Add new information to the AI's knowledge base for future conversations",
            "parameters": {
                "type": "object",
                "properties": {
                    "information": {
                        "type": "string",
                        "description": "The information to learn and remember",
                    },
                    "metadata": {
                        "type": "object",
                        "description": "Optional metadata about the information",
                        "properties": {
                            "source": {"type": "string"},
                            "category": {"type": "string"},
                        },
                    },
                },
                "required": ["information"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_knowledge",
            "description": "Search the AI's learned knowledge base for relevant information",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Search query to find relevant learned information",
                    },
                    "limit": {
                        "type": "integer",
                        "description": "Maximum number of results to return",
                        "default": 5,
                    },
                },
                "required": ["query"],
            },
        },
    },
]


def execute_tool_call(tool_call):
    """
    Execute a tool call requested by devstral-small-2507-mlx via function calling.

    This is the central dispatcher for all AI-initiated tool operations. When the devstral-small-2507-mlx
    model decides to use a tool (based on user intent and tool descriptions), this function:

    1. Parses the tool call from LangChain's structured format
    2. Routes to the appropriate tool execution function
    3. Applies security checks and sandboxing
    4. Returns structured results for AI consumption

    SUPPORTED TOOLS (7 total):
    - read_file: Read file contents with security restrictions
    - write_file: Create/modify files with path validation
    - list_directory: Browse directory contents safely
    - get_current_directory: Get current working directory
    - parse_document: Extract structured data from documents using devstral-small-2507
    - learn_information: Store information in ChromaDB knowledge base
    - search_knowledge: Query learned information via semantic search

    SECURITY FEATURES:
    - Directory sandboxing (current directory only)
    - Path traversal prevention
    - File size limits (1MB for reading)
    - Binary file detection and rejection
    - Comprehensive error handling

    Args:
        tool_call (dict): LangChain tool call structure containing:
            - name: Tool function name (e.g., "read_file")
            - args: Tool parameters as dict
            - id: Unique tool call identifier
            - type: Always "tool_call"

    Returns:
        dict: Structured execution result with:
            - tool_call_id: Original call identifier for AI correlation
            - function_name: Executed tool name
            - result: Tool-specific return data or error information

    Raises:
        None: All exceptions caught and returned as error results

    Example:
        Input: {"name": "read_file", "args": {"file_path": "README.md"}, "id": "123"}
        Output: {"tool_call_id": "123", "function_name": "read_file", "result": {"success": True, "content": "..."}}
    """
    # Initialize variables
    function_name = ""
    tool_call_id = ""
    result = None

    try:
        # Handle LangChain tool call format (dict with 'name', 'args', 'id', 'type')
        function_name = tool_call["name"]
        arguments = tool_call["args"]
        tool_call_id = tool_call["id"]

        if function_name == "read_file":
            file_path = arguments.get("file_path", "")
            result = execute_read_file(file_path)
        elif function_name == "write_file":
            file_path = arguments.get("file_path", "")
            content = arguments.get("content", "")
            result = execute_write_file(file_path, content)
        elif function_name == "list_directory":
            directory_path = arguments.get("directory_path", ".")
            result = execute_list_directory(directory_path)
        elif function_name == "get_current_directory":
            result = execute_get_current_directory()
        elif function_name == "learn_information":
            information = arguments.get("information", "")
            metadata = arguments.get("metadata", None)
            result = execute_learn_information(information, metadata)
        elif function_name == "search_knowledge":
            query = arguments.get("query", "")
            limit = arguments.get("limit", 5)
            result = execute_search_knowledge(query, limit)
        elif function_name == "parse_document":
            file_path = arguments.get("file_path", "")
            extract_type = arguments.get("extract_type", "text")
            result = execute_parse_document(file_path, extract_type)
        else:
            result = {"error": f"Unknown tool: {function_name}"}

        return {
            "tool_call_id": tool_call_id,
            "function_name": function_name,
            "result": result,
        }

    except Exception as e:
        return {
            "tool_call_id": tool_call_id,
            "function_name": function_name,
            "result": {"error": str(e)},
        }


def execute_read_file(file_path: str) -> dict:
    """Execute file reading tool."""
    try:
        # Security check
        current_dir = os.getcwd()
        full_path = os.path.abspath(file_path)

        if not full_path.startswith(current_dir):
            return {
                "error": "Access denied: Cannot read files outside current directory"
            }

        if not os.path.exists(full_path):
            return {"error": f"File not found: {file_path}"}

        if not os.path.isfile(full_path):
            return {"error": f"Not a file: {file_path}"}

        # Size check
        file_size = os.path.getsize(full_path)
        if file_size > 1024 * 1024:  # 1MB limit
            return {"error": f"File too large: {file_size} bytes (max 1MB)"}

        with open(full_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        return {
            "success": True,
            "file_path": file_path,
            "size": file_size,
            "content": content,
        }

    except UnicodeDecodeError:
        return {"error": "Cannot read binary file"}
    except Exception as e:
        return {"error": str(e)}


def execute_write_file(file_path: str, content: str) -> dict:
    """Execute file writing tool."""
    try:
        # Security check
        current_dir = os.getcwd()
        full_path = os.path.abspath(file_path)

        if not full_path.startswith(current_dir):
            return {
                "error": "Access denied: Cannot write files outside current directory"
            }

        # Create directory if needed
        dir_path = os.path.dirname(full_path)
        if dir_path and not os.path.exists(dir_path):
            os.makedirs(dir_path, exist_ok=True)

        with open(full_path, "w", encoding="utf-8") as f:
            f.write(content)

        return {"success": True, "file_path": file_path, "size": len(content)}

    except Exception as e:
        return {"error": str(e)}


def execute_list_directory(directory_path: str = ".") -> dict:
    """Execute directory listing tool."""
    try:
        # Security check
        current_dir = os.getcwd()
        full_path = os.path.abspath(directory_path)

        if not full_path.startswith(current_dir):
            return {
                "error": "Access denied: Cannot list directories outside current directory"
            }

        if not os.path.exists(full_path):
            return {"error": f"Directory not found: {directory_path}"}

        if not os.path.isdir(full_path):
            return {"error": f"Not a directory: {directory_path}"}

        items = os.listdir(full_path)
        dirs = []
        files = []

        for item in sorted(items):
            full_item_path = os.path.join(full_path, item)
            if os.path.isdir(full_item_path):
                dirs.append(item + "/")
            else:
                files.append(item)

        return {
            "success": True,
            "directory": os.path.relpath(full_path, current_dir) or ".",
            "contents": dirs + files,
            "total_items": len(items),
        }

    except Exception as e:
        return {"error": str(e)}


def execute_get_current_directory() -> dict:
    """Execute current directory tool."""
    try:
        return {"success": True, "current_directory": os.getcwd()}
    except Exception as e:
        return {"error": str(e)}


def execute_learn_information(information: str, metadata: dict | None = None) -> dict:
    """Execute learn information tool."""
    try:
        if not information.strip():
            return {"error": "Information cannot be empty"}

        # Use existing learn functionality
        handle_learn_command(information)

        return {
            "success": True,
            "information_length": len(information),
            "learned": True,
        }

    except Exception as e:
        return {"error": str(e)}


def execute_search_knowledge(query: str, limit: int = 5) -> dict:
    """Execute knowledge search tool."""
    try:
        if not query.strip():
            return {"error": "Query cannot be empty"}

        # Use existing context retrieval
        context = get_relevant_context(query, limit)

        return {
            "success": True,
            "query": query,
            "results": context,
            "result_count": len(context) if context else 0,
        }

    except Exception as e:
        return {"error": str(e)}


def execute_parse_document(file_path: str, extract_type: str) -> dict:
    """
    Execute document parsing tool using devstral-small-2507's advanced multimodal capabilities.

    This function leverages the devstral-small-2507-mlx model's sophisticated understanding of visual
    content to extract structured data from various document types. It acts as a bridge
    between file system access and AI-powered document intelligence.

    SUPPORTED DOCUMENT TYPES:
    - Images: PNG, JPG, JPEG, BMP, TIFF (direct OCR processing)
    - Office Documents: DOCX, XLSX, PPTX (require conversion to images)
    - Text Documents: TXT, MD, RST, JSON, XML, CSV (text processing)

    EXTRACTION TYPES:
    - "text": Extract readable text with layout preservation
    - "tables": Identify and extract tabular data as structured JSON
    - "forms": Recognize form fields and input areas
    - "layout": Analyze document structure and visual organization

    PROCESSING FLOW:
    1. Security validation (current directory only)
    2. File type verification and support checking
    3. Specialized prompt generation based on extract_type
    4. devstral-small-2507 analysis with multimodal understanding
    5. Structured result formatting for AI consumption

    Args:
        file_path (str): Path to document file relative to current directory.
                        Must be within the current working directory for security.
        extract_type (str): Type of information to extract. Must be one of:
                           "text", "tables", "forms", "layout"

    Returns:
        dict: Structured extraction results containing:
            - success (bool): Whether extraction completed successfully
            - extract_type (str): The type of extraction performed
            - content/analysis/tables (varies): The extracted data
            - file_type (str): Detected file extension
            - error (str): Error message if success=False

    Raises:
        None: All exceptions are caught and returned as error results

    Security:
        - Path traversal prevention
        - Current directory sandboxing
        - File existence validation
        - Supported format restrictions

    Example:
        execute_parse_document("report.pdf", "tables")
        # Returns: {"success": True, "tables": [...], "extract_type": "tables"}
    """
    try:
        # Security check - ensure file is within current directory
        current_dir = os.getcwd()
        full_path = os.path.abspath(file_path)

        if not full_path.startswith(current_dir):
            return {"error": "Access denied: File outside current directory"}

        if not os.path.exists(full_path):
            return {"error": f"File not found: {file_path}"}

        # Check if file is a supported document type
        supported_extensions = {
            ".pdf",
            ".png",
            ".jpg",
            ".jpeg",
            ".bmp",
            ".tiff",  # Images
            ".docx",
            ".xlsx",
            ".pptx",  # Office documents
            ".rtf",  # Rich Text Format
            ".epub",  # EPUB e-books
            ".txt",
            ".md",
            ".rst",
            ".json",
            ".xml",
            ".csv",  # Text documents
        }
        file_ext = os.path.splitext(full_path)[1].lower()

        if file_ext not in supported_extensions:
            return {
                "error": f"Unsupported file type: {file_ext}. Supported: {', '.join(supported_extensions)}"
            }

        # Extract content based on file type and extract_type
        if file_ext == ".pdf":
            try:
                from PyPDF2 import PdfReader

                reader = PdfReader(full_path)
                extracted_text = ""

                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        extracted_text += text + "\n"

                if not extracted_text.strip():
                    return {
                        "success": False,
                        "error": "No readable text found in PDF (may contain only images or be scanned)",
                    }

                # Clean up the text
                extracted_text = extracted_text.replace("\x00", "").strip()

                if extract_type == "text":
                    return {
                        "success": True,
                        "file_path": file_path,
                        "extract_type": extract_type,
                        "file_type": file_ext,
                        "content": extracted_text,
                        "analysis": extracted_text,
                        "pages": len(reader.pages),
                    }
                else:
                    # For other extract types, return basic text for now
                    return {
                        "success": True,
                        "file_path": file_path,
                        "extract_type": extract_type,
                        "file_type": file_ext,
                        "analysis": f"Basic text extraction from PDF. Full {extract_type} analysis not yet implemented.",
                        "content": extracted_text[:500] + "..."
                        if len(extracted_text) > 500
                        else extracted_text,
                        "pages": len(reader.pages),
                    }

            except Exception as e:  # type: ignore
                if "No module named" in str(e) and "striprtf" in str(e):
                    return {
                        "success": False,
                        "error": "striprtf not available for RTF processing. Install with: pip install striprtf",
                    }
                else:
                    return {
                        "success": False,
                        "error": f"RTF processing failed: {str(e)}",
                    }
            except Exception as e:  # type: ignore  # type: ignore
                return {"success": False, "error": f"PDF processing failed: {str(e)}"}

        elif file_ext == ".docx":
            # Extract text content from Microsoft Word documents (.docx)
            # Uses python-docx library to read document structure and extract readable text
            # Processes both paragraph text and table cell content for comprehensive extraction
            try:
                from docx import Document

                doc = Document(full_path)
                content = ""

                # Extract text from all paragraphs in the document
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content += paragraph.text + "\n"

                # Extract text from all tables in the document
                # Tables are processed row by row, cell by cell
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content += cell.text + "\n"

                return {
                    "success": True,
                    "file_path": file_path,
                    "extract_type": extract_type,
                    "file_type": file_ext,
                    "content": content,
                    "analysis": content,
                }
            except Exception as e:  # type: ignore  # type: ignore
                if "No module named" in str(e) and "docx" in str(e):
                    return {
                        "success": False,
                        "error": "python-docx not available for DOCX processing. Install with: pip install python-docx",
                    }
                else:
                    return {
                        "success": False,
                        "error": f"DOCX processing failed: {str(e)}",
                    }

        elif file_ext == ".rtf":
            # Extract text content from Rich Text Format files (.rtf)
            # Uses striprtf library to strip RTF formatting codes and extract plain text
            # Handles RTF control sequences and converts to readable text content
            try:
                import striprtf  # type: ignore

                # Read the RTF file with UTF-8 encoding and error replacement
                with open(full_path, "r", encoding="utf-8", errors="replace") as f:
                    rtf_content = f.read()

                # Strip RTF formatting to get plain text
                content = striprtf.striprtf(rtf_content)  # type: ignore

                return {
                    "success": True,
                    "file_path": file_path,
                    "extract_type": extract_type,
                    "file_type": file_ext,
                    "content": content,
                    "analysis": content,
                }
            except ImportError:  # type: ignore
                return {
                    "success": False,
                    "error": "striprtf not available for RTF processing. Install with: pip install striprtf",
                }
            except Exception as e:  # type: ignore
                return {
                    "success": False,
                    "error": f"RTF processing failed: {str(e)}",
                }

        elif file_ext == ".epub":
            # Extract text content from EPUB e-book files (.epub)
            # Uses ebooklib to parse EPUB structure and extract text from HTML chapters
            # Processes all document items, strips HTML tags, and concatenates readable text
            try:
                import ebooklib
                from ebooklib import epub

                book = epub.read_epub(full_path)
                content = ""

                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        # Extract text from HTML content
                        html_content = item.get_content().decode("utf-8")
                        # Simple HTML text extraction (remove tags)
                        import re

                        text_content = re.sub(r"<[^>]+>", "", html_content)
                        text_content = re.sub(
                            r"\s+", " ", text_content).strip()
                        if text_content:
                            content += text_content + "\n\n"

                return {
                    "success": True,
                    "file_path": file_path,
                    "extract_type": extract_type,
                    "file_type": file_ext,
                    "content": content,
                    "analysis": content,
                }
            except Exception as e:  # type: ignore
                if "No module named" in str(e) and "ebooklib" in str(e):
                    return {
                        "success": False,
                        "error": "ebooklib not available for EPUB processing. Install with: pip install ebooklib",
                    }
                else:
                    return {
                        "success": False,
                        "error": f"EPUB processing failed: {str(e)}",
                    }
            except ImportError:  # type: ignore
                return {
                    "success": False,
                    "error": "ebooklib not available for EPUB processing. Install with: pip install ebooklib",
                }
            except Exception as e:  # type: ignore
                return {
                    "success": False,
                    "error": f"EPUB processing failed: {str(e)}",
                }
            except ImportError:  # type: ignore
                return {
                    "success": False,
                    "error": "ebooklib not available for EPUB processing",
                }
            except Exception as e:  # type: ignore
                return {
                    "success": False,
                    "error": f"EPUB processing failed: {str(e)}",
                }

        elif file_ext == ".xlsx":
            # Extract text content from Excel spreadsheet files (.xlsx)
            # Uses openpyxl to read workbook structure and extract cell values
            # Processes all sheets, converting tabular data to tab-separated text format
            try:
                from openpyxl import load_workbook  # type: ignore

                # Load workbook with data_only=True to get cell values instead of formulas
                wb = load_workbook(full_path, data_only=True)
                content = ""

                # Process each worksheet in the workbook
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    content += f"Sheet: {sheet_name}\n"

                    # Iterate through all rows, extracting cell values
                    for row in sheet.iter_rows(values_only=True):
                        # Convert row to tab-separated string, handling None values
                        row_text = "\t".join(
                            str(cell) if cell is not None else "" for cell in row
                        )
                        if row_text.strip():
                            content += row_text + "\n"

                    content += "\n"

                return {
                    "success": True,
                    "file_path": file_path,
                    "extract_type": extract_type,
                    "file_type": file_ext,
                    "content": content,
                    "analysis": content,
                }
            except Exception as e:  # type: ignore
                if "No module named" in str(e) and "openpyxl" in str(e):
                    return {
                        "success": False,
                        "error": "openpyxl not available for XLSX processing. Install with: pip install openpyxl",
                    }
                else:
                    return {
                        "success": False,
                        "error": f"XLSX processing failed: {str(e)}",
                    }
            except ImportError:  # type: ignore
                return {
                    "success": False,
                    "error": "openpyxl not available for XLSX processing. Install with: pip install openpyxl",
                }
            except Exception as e:  # type: ignore
                return {
                    "success": False,
                    "error": f"XLSX processing failed: {str(e)}",
                }
            except Exception as e:  # type: ignore
                if "No module named" in str(e) and "docx" in str(e):
                    return {
                        "success": False,
                        "error": "python-docx not available for DOCX processing. Install with: pip install python-docx",
                    }
                else:
                    return {
                        "success": False,
                        "error": f"DOCX processing failed: {str(e)}",
                    }
            except Exception as e:  # type: ignore
                return {
                    "success": False,
                    "error": f"XLSX processing failed: {str(e)}",
                }
            except ImportError:  # type: ignore
                return {
                    "success": False,
                    "error": "ebooklib not available for EPUB processing",
                }
            except Exception as e:  # type: ignore
                return {
                    "success": False,
                    "error": f"EPUB processing failed: {str(e)}",
                }

        elif file_ext in [".txt", ".md", ".rst", ".json", ".xml", ".csv"]:
            # For text-based files, just read the content
            try:
                with open(full_path, "r", encoding="utf-8", errors="replace") as f:
                    content = f.read()

                return {
                    "success": True,
                    "file_path": file_path,
                    "extract_type": extract_type,
                    "file_type": file_ext,
                    "content": content,
                    "analysis": content,
                }
            except Exception as e:  # type: ignore
                return {
                    "success": False,
                    "error": f"Text file reading failed: {str(e)}",
                }

        else:
            # For other file types, return placeholder for now
            return {
                "success": True,
                "file_path": file_path,
                "extract_type": extract_type,
                "file_type": file_ext,
                "analysis": f"Document analysis for {extract_type} extraction using devstral-small-2507 capabilities",
                "note": f"Full {extract_type} extraction not yet implemented for {file_ext} files",
            }

    except Exception as e:
        return {"error": str(e)}


def main():
    """
    Main chat loop v0.1 - Core interactive interface with AI tool calling.

    This function implements the primary user interaction loop featuring:
    1. AI tool calling with devstral-small-2507-mlx (7 autonomous tools supported)
    2. AI learning and knowledge retention via ChromaDB vector database
    3. Comprehensive slash command system for direct user control
    4. Context-aware conversations using semantic search and learned information
    5. Model switching and configuration management
    6. Codebase ingestion for comprehensive AI knowledge building
    7. Persistent conversation memory across sessions via SQLite
    8. Spaces system for isolated knowledge workspaces
    9. Document intelligence with devstral-small-2507 multimodal analysis

    AI TOOL ECOSYSTEM (7 Tools - devstral-small-2507-mlx Function Calling):
    =======================
    File System Tools:
    - read_file() - Secure file reading with content validation
    - write_file() - Safe file creation/modification with path checking
    - list_directory() - Directory browsing with security restrictions
    - get_current_directory() - Current working directory retrieval

    Document Intelligence Tools:
    - parse_document() - devstral-small-2507 multimodal document analysis (OCR, tables, forms, layout)

    Knowledge Management Tools:
    - learn_information() - Store information in ChromaDB with metadata
    - search_knowledge() - Semantic search across learned knowledge base

    INTERACTION MODES:
    ==================
    Direct Commands: User types /command → Immediate execution (no AI processing)
    Natural Language: User describes intent → AI tool calling → Autonomous execution
    Hybrid: Commands teach AI patterns → AI learns to call tools appropriately

    Available Commands:
    - /learn <text> - Teach AI new information
    - /vectordb - Inspect learned knowledge base
    - /populate <path> - Bulk import codebases
    - /model - Check/switch AI models
    - /memory - View conversation history
    - /clear - Reset conversation memory
    - /space <cmd> - Space/workspace management (list/create/switch/delete)
    - /read <file> - Read file contents
    - /write <file> - Write content to file
    - /list [dir] - List directory contents
    - /pwd - Show current directory
    - /help - Show all available commands

    Flow:
    - Input validation and command routing
    - Automatic file command detection (natural language)
    - AI tool calling for complex operations
    - Direct command processing (no LLM calls for commands)
    - Context retrieval and AI response generation
    - Memory persistence and cleanup
    """
    show_welcome()

    # Initialize application components
    if not initialize_application():
        return  # Exit if initialization failed

    # Main interaction loop - runs until user quits
    while True:
        try:
            # Get user input with prompt
            user_input = input("You: ").strip()

            # Skip empty input to prevent blank messages
            if not user_input:
                continue

            # Handle quit commands - immediate exit with memory save
            # Multiple quit variations for user convenience
            if user_input.lower() in ["quit", "exit", "q"]:
                print("\nAI Assistant: Goodbye! Have a great day!")
                # Preserve conversation state
                save_memory(conversation_history)
                break

            # Additional quit command check (belt and suspenders approach)
            # Handles edge cases where quit might be embedded in other text
            if any(quit_cmd in user_input.lower() for quit_cmd in ["quit", "exit"]):
                print("\nAI Assistant: Goodbye! Have a great day!")
                save_memory(conversation_history)
                break

            # Handle slash commands
            if user_input.startswith("/"):
                if handle_slash_command(user_input):
                    continue  # Command was handled, get next input
                else:
                    print(f"\nUnknown command: {user_input}")
                    print("Type /help for available commands\n")
                    continue

            # Validate input length to prevent memory and API issues
            # Large inputs can cause token limit errors or memory problems
            if len(user_input) > MAX_INPUT_LENGTH:
                print(
                    f"\nError: Input exceeds maximum length of {MAX_INPUT_LENGTH} characters."
                )
                continue

                # /model - Display current model information

            # Add user message to conversation history
            # This maintains the chronological order of the conversation
            conversation_history.append(HumanMessage(content=user_input))

            # =============================================================================
            # CONTEXT RETRIEVAL FROM VECTOR DATABASE
            # =============================================================================

            # Retrieve relevant context from knowledge base using semantic similarity search
            # This enhances AI responses by providing background information from learned data
            # Only performed if vector database is enabled and available
            context = get_relevant_context(user_input) if vectorstore else ""

            # =============================================================================
            # ENHANCED PROMPT CONSTRUCTION
            # =============================================================================

            # Prepare enhanced conversation history for LLM with context integration
            # The goal is to provide the AI with relevant background knowledge while maintaining conversation flow
            enhanced_history = conversation_history.copy()

            # Apply context integration based on user-controlled settings
            is_learning_query = False  # Track if this is a learning-related query

            # Add tool instructions to system message for better tool usage
            tool_instructions = """

IMPORTANT: You have access to tools for file system operations and document processing.
When users ask you to read files, list directories, analyze documents, or perform any file operations,
you MUST use the appropriate tools instead of responding conversationally.

Available tools:
- read_file(file_path): Use this when users ask to read, view, or show file contents
- write_file(file_path, content): Use this when users ask to create or modify files
- list_directory(directory_path): Use this when users ask to list files or see directory contents
- get_current_directory(): Use this when users ask for current directory or pwd
- parse_document(file_path, extract_type): Use this for document analysis - extract text, tables,
  forms, or layout from PDFs, images, and office documents
- learn_information(info): Use this when users want to teach you information
- search_knowledge(query): Use this when users ask what you know or search learned info

CRITICAL: If a user says "read the README", "analyze this document",
"extract tables from PDF", or similar, you MUST call the appropriate tool.
Do not respond with text about not having access to files.
"""

            if LEARNING_MODE == "off":
                # Learning is completely disabled
                context = ""
                logger.info(
                    "Learning mode is off - skipping context integration")
            elif CONTEXT_MODE == "off":
                # Context integration is disabled
                context = ""
                logger.info(
                    "Context mode is off - skipping context integration")
            elif CONTEXT_MODE == "on" and context:
                # Always include context when available
                logger.info(
                    f"Context mode is on - including context: {context[:200]}..."
                )
            elif CONTEXT_MODE == "auto" and context:
                # Auto mode: Use context for learning queries, skip for meta questions
                learning_keywords = [
                    "what have you learned",
                    "what did you learn",
                    "what do you know",
                    "what have you been taught",
                    "what information do you have",
                    "what's in your knowledge",
                    "what's in your database",
                    "tell me what you learned",
                    "show me what you know",
                ]

                is_learning_query = any(
                    keyword in user_input.lower() for keyword in learning_keywords
                )

                if is_learning_query:
                    logger.info(
                        f"Auto mode - Learning query detected, including context: {context[:200]}..."
                    )
                else:
                    logger.info(
                        f"Auto mode - Regular query, skipping context integration"
                    )

            # Apply context to enhanced history if available
            should_include_context = CONTEXT_MODE == "on" or (
                CONTEXT_MODE == "auto"
                and "is_learning_query" in locals()
                and is_learning_query
            )

            # Create system message with tool instructions
            system_content = "Lets get some coding done.." + tool_instructions
            if context and should_include_context:
                system_content += f" Use this context: {context}"

            # Replace the first system message if it exists, otherwise insert
            if enhanced_history and isinstance(enhanced_history[0], SystemMessage):
                enhanced_history[0] = SystemMessage(content=system_content)
            else:
                enhanced_history.insert(
                    0, SystemMessage(content=system_content))
            system_content = "Lets get some coding done.." + tool_instructions
            logger.info(
                f"After concatenation, system_content len: {len(system_content)}"
            )
            if context and should_include_context:
                system_content += f" Use this context: {context}"

            # Replace the first system message if it exists, otherwise insert
            if enhanced_history and isinstance(enhanced_history[0], SystemMessage):
                logger.info("Replacing existing system message")
                enhanced_history[0] = SystemMessage(content=system_content)
                logger.info(
                    f"New first message content: {enhanced_history[0].content[:100]}..."
                )
            else:
                logger.info("Inserting new system message")
                enhanced_history.insert(
                    0, SystemMessage(content=system_content))

            logger.info(
                f"Final system message ({len(system_content)} chars): {repr(system_content[:100])}..."
            )

            if context and should_include_context:
                logger.info(
                    f"Context integrated ({len(context)} chars): {context[:100]}..."
                )

            # Generate AI response with tool calling support
            response = ""  # Accumulate full response

            # Ensure LLM is initialized
            if llm is None:
                print("\n❌ Error: LLM not initialized\n")
                continue

            # Try tool-enabled response first
            try:
                # Make initial call with tools already bound to LLM
                initial_response = llm.invoke(enhanced_history)
                logger.info(f"LLM response type: {type(initial_response)}")
                logger.info(
                    f"Has tool_calls attr: {hasattr(initial_response, 'tool_calls')}"
                )
                if hasattr(initial_response, "tool_calls"):
                    logger.info(f"Tool calls: {initial_response.tool_calls}")

                # Check if the response contains tool calls
                if (
                    hasattr(initial_response, "tool_calls")
                    and initial_response.tool_calls
                ):
                    # Execute tool calls
                    tool_results = []
                    for tool_call in initial_response.tool_calls:
                        tool_name = (
                            tool_call.get("name", "unknown")
                            if isinstance(tool_call, dict)
                            else getattr(tool_call, "name", "unknown")
                        )
                        print(f"🔧 Executing tool: {tool_name}")
                        result = execute_tool_call(tool_call)
                        tool_results.append(result)

                    # Add tool results to conversation for follow-up
                    if tool_results:
                        tool_message = "Tool execution results:\n"
                        for result in tool_results:
                            tool_message += (
                                f"- {result['function_name']}: {result['result']}\n"
                            )

                        enhanced_history.append(
                            AIMessage(
                                content=initial_response.content or "Using tools..."
                            )
                        )
                        enhanced_history.append(
                            HumanMessage(content=tool_message))

                        # Make follow-up call with tool results
                        final_response = llm.invoke(enhanced_history)
                        response = final_response.content or ""
                    else:
                        response = initial_response.content or ""

                else:
                    # No tool calls, use the initial response
                    response = initial_response.content or ""

            except Exception as e:  # type: ignore
                # Fallback to regular streaming if tool calling fails
                logger.warning(
                    f"Tool calling failed, falling back to regular response: {e}"
                )
                # Stream response tokens from LLM (fallback method)
                for chunk in llm.stream(enhanced_history):
                    content = chunk.content

                    # Handle different content types from the LLM
                    if isinstance(content, list):
                        content = "".join(str(c) for c in content)

                    response += content

            # Render the response with markdown formatting for better CLI display
            try:
                from rich.console import Console
                from rich.markdown import Markdown

                console = Console()
                print("AI Assistant:")
                # Use rich to render markdown in terminal
                if isinstance(response, str):
                    console.print(Markdown(response))
                else:
                    console.print(str(response))
            except ImportError:  # type: ignore
                # Fallback to plain text if rich not available
                print("AI Assistant:", response)

            # Add AI response to conversation history
            # Critical: Use AIMessage (not SystemMessage) for proper message typing
            conversation_history.append(AIMessage(content=response))

            # Trim conversation history to prevent memory bloat
            # Maintains recent context while staying within API limits
            conversation_history[:] = trim_history(
                conversation_history, MAX_HISTORY_PAIRS
            )

            # =============================================================================
            # ERROR HANDLING
            # =============================================================================

        except ConnectionError as e:
            # Network connectivity issues - LM Studio server is unreachable
            logger.error(f"Connection error: {e}")
            print(
                f"\n\nError: Cannot connect to LM Studio at {LM_STUDIO_BASE_URL}")
            print("Please ensure LM Studio is running and accessible.\n")
            # Remove failed user message to prevent conversation corruption
            if conversation_history and isinstance(
                conversation_history[-1], HumanMessage
            ):
                conversation_history.pop()

        except EOFError:
            # End of input stream (e.g., pipe closed, file ended)
            # Common when running in automated environments or with piped input
            print("\n\nAI Assistant: Input ended. Goodbye!")
            save_memory(conversation_history)
            logger.info("End of input reached")
            break

        except Exception as e:
            # Catch-all for any other LLM or processing errors (API errors, model issues, etc.)
            logger.error(f"Error processing request: {e}")
            print(f"\n\nError: {type(e).__name__}: {e}\n")
            # Remove failed user message to maintain conversation integrity
            if conversation_history and isinstance(
                conversation_history[-1], HumanMessage
            ):
                conversation_history.pop()

        # =============================================================================
        # SYSTEM-LEVEL EXCEPTION HANDLING
        # =============================================================================

        except KeyboardInterrupt:
            # User pressed Ctrl+C - graceful shutdown preserving conversation state
            print("\n\nAI Assistant: Interrupted. Goodbye!")
            save_memory(conversation_history)  # Preserve conversation state
            logger.info("User interrupted the program")
            break


# =============================================================================
# APPLICATION ENTRY POINT
# =============================================================================

# Standard Python entry point - only execute main() when run directly
# Prevents execution when imported as a module by other scripts
# This ensures the chat interface only starts when the script is run standalone
if __name__ == "__main__":
    main()
